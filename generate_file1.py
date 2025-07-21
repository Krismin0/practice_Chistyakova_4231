from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import docx

RU_MONTHS = {
    'January': 'января', 'February': 'февраля', 'March': 'марта',
    'April': 'апреля', 'May': 'мая', 'June': 'июня',
    'July': 'июля', 'August': 'августа', 'September': 'сентября',
    'October': 'октября', 'November': 'ноября', 'December': 'декабря'
}


def set_document_styles(doc):
    """Устанавливает основные стили документа"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(0)
    style.element.xpath(".//w:pPr")[0].append(
        docx.oxml.parse_xml(
            r'<w:contextualSpacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    )


def format_date_russian(date_str):
    """Форматирует дату в русский формат (день месяц)"""
    try:
        date_obj = datetime.strptime(date_str.strip(), "%Y-%m-%d")
        eng_month = date_obj.strftime("%B")
        ru_month = RU_MONTHS.get(eng_month, eng_month)
        return f"{date_obj.day} {ru_month}"
    except Exception:
        return date_str


def process_contributions(contributions):
    """Обрабатывает данные из JSON"""
    processed = []
    for item in contributions:
        persons = item.get("persons", [{}])
        speaker = persons[0] if persons else {}

        group = ""
        for field in item.get("custom_fields", []):
            if field.get("name") == "Номер группы основного автора (докладчика)":
                group = field.get("value", "")
                break

        start_dt = item.get("start_dt", "")
        date = start_dt.split("T")[0] if start_dt else ""
        time = start_dt.split("T")[1][:5] if start_dt else ""

        processed.append({
            "session": item.get("session", {}).get("friendly_id", "1"),
            "date": date,
            "time": time,
            "aud": item.get("room_name", ""),
            "full_name": speaker.get("full_name", ""),
            "group": group,
            "topic": item.get("title", "")
        })
    return processed


def add_conference_header(doc, conference_data):
    """Добавляет шапку программы конференции"""
    # Заголовок программы
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Программа {conference_data['number']} МСНК ГУАП")
    run.bold = True
    run.italic = True
    run.font.size = Pt(14)

    # Подзаголовок
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("по кафедре № 43 компьютерных технологий и программной инженерии")
    run.bold = True
    run.italic = True
    run.font.size = Pt(14)
    subtitle.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph("")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15

    # Секция кафедры (специальное форматирование)
    p_section = doc.add_paragraph()
    p_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_section.paragraph_format.first_line_indent = Cm(1.27)
    run = p_section.add_run("Секция каф.43. «компьютерных технологий и программной инженерии»")
    run.bold = True
    run.italic = True
    run.font.size = Pt(12)
    p_section.paragraph_format.space_after = Pt(16)

    # Руководство секции (не жирное, не курсив, отступ 2 см)
    leader_info = [
        f"Научный руководитель секции – {conference_data['head']}",
        conference_data.get('head_title', ''),
        f"Зам. научного руководителя секции – {conference_data['deputy']}",
        conference_data.get('deputy_title', ''),
        f"Секретарь – {conference_data['secretary']}",
        conference_data.get('secretary_title', '')
    ]

    for line in leader_info:
        if line.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(2)
            p.add_run(line)
            p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph("")


def add_session(doc, session_num, session_data, is_first_session=False):
    """Добавляет информацию о заседании"""
    date_str = format_date_russian(session_data[0]["date"])
    time = session_data[0]["time"]
    aud = session_data[0]["aud"]

    # Заголовок заседания
    p1 = doc.add_paragraph()
    run = p1.add_run(f"Заседание {session_num}.")
    run.bold = True
    run.font.size = Pt(14)
    p1.paragraph_format.space_after = Pt(6)

    # Дата и место
    p2 = doc.add_paragraph()
    run = p2.add_run(f"{date_str}, {time}, ауд. {aud}.")
    run.bold = True
    run.font.size = Pt(12)
    p2.paragraph_format.space_after = Pt(6)

    # Служебная фраза (только для первого заседания)
    if is_first_session:
        p3 = doc.add_paragraph()
        run = p3.add_run("По решению руководителя секции порядок следования докладов может быть изменен.")
        run.italic = True
        run.font.size = Pt(14)
        p3.paragraph_format.space_after = Pt(0)
        p = doc.add_paragraph("")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15

    # Участники (шрифт 14 pt)
    for i, item in enumerate(session_data, 1):
        if not item["full_name"] or not item["topic"]:
            continue

        # Имя и группа
        para = doc.add_paragraph(style='List Number')
        para.paragraph_format.left_indent = Cm(1.25)
        para.paragraph_format.first_line_indent = Cm(-0.75)
        para.paragraph_format.space_after = Pt(0)

        full = f"{item['full_name']}"
        if item["group"]:
            full += f", группа {item['group']}"

        run = para.add_run(full)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)  # Шрифт 14 для списка студентов

        # Тема доклада (шрифт 14 pt)
        title = doc.add_paragraph()
        title.paragraph_format.left_indent = Cm(1.25)
        title.paragraph_format.space_after = Pt(12)
        run2 = title.add_run(item["topic"])
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(14)  # Шрифт 14 для тем докладов


def generate_conference_program(conference_data, contributions, output_path):
    """Генерирует файл программы конференции"""
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    set_document_styles(doc)

    processed_entries = process_contributions(contributions)
    add_conference_header(doc, conference_data)

    sessions = {}
    for entry in processed_entries:
        session = entry["session"]
        if session not in sessions:
            sessions[session] = []
        sessions[session].append(entry)

    # Сортируем заседания и добавляем их в документ
    sorted_sessions = sorted(sessions.items(), key=lambda x: int(x[0]))
    for i, (session_num, session_data) in enumerate(sorted_sessions):
        add_session(doc, session_num, session_data, is_first_session=(i == 0))

    doc.save(output_path)