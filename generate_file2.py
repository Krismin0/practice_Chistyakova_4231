from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

RU_MONTHS = {
    'January': 'января', 'February': 'февраля', 'March': 'марта',
    'April': 'апреля', 'May': 'мая', 'June': 'июня',
    'July': 'июля', 'August': 'августа', 'September': 'сентября',
    'October': 'октября', 'November': 'ноября', 'December': 'декабря'
}


def set_document_styles(doc):
    """Устанавливает основные стили документа"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(0)


def format_date_russian(date_str):
    """Форматирует дату в русский формат (день месяц год)"""
    try:
        date_obj = datetime.strptime(date_str.strip(), "%Y-%m-%d")
        eng_month = date_obj.strftime("%B")
        ru_month = RU_MONTHS.get(eng_month, eng_month)
        return f"{date_obj.day} {ru_month} {date_obj.year} г."
    except Exception:
        return date_str


def get_status_from_group(group_number):
    """Определяет статус (магистрант/бакалавр) по номеру группы"""
    if not group_number:
        return ""

    last_char = group_number[-1].upper()
    if last_char in ('M', 'М'):
        return "магистр"
    return "студент"


def process_contributions(contributions):
    """Обрабатывает сырые данные из JSON для отчета"""
    processed = []
    for item in contributions:
        persons = item.get("persons", [{}])
        speaker = persons[0] if persons else {}

        group = ""
        for field in item.get("custom_fields", []):
            if field.get("name") == "Номер группы основного автора (докладчика)":
                group = field.get("value", "")
                break

        start_dt = item.get("start_dt", "")
        date = start_dt.split("T")[0] if start_dt else ""
        time = start_dt.split("T")[1][:5] if start_dt else ""

        processed.append({
            "session": item.get("session", {}).get("friendly_id", "1"),
            "date": date,
            "time": time,
            "aud": item.get("room_name", ""),
            "full_name": speaker.get("full_name", ""),
            "group": group,
            "topic": item.get("title", ""),
            "status": get_status_from_group(group)
        })
    return processed


def set_cell_format(cell, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Устанавливает форматирование для ячейки таблицы"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.font.name = 'Times New Roman'
    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTH


def set_table_borders(table):
    """Устанавливает границы таблицы толщиной 1 pt"""
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblBorders = tblPr.first_child_found_in('w:tblBorders')
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    for border in borders:
        border_element = OxmlElement(f'w:{border}')
        border_element.set(qn('w:val'), 'single')
        border_element.set(qn('w:sz'), '8')
        border_element.set(qn('w:space'), '0')
        border_element.set(qn('w:color'), '000000')
        tblBorders.append(border_element)


def add_session_table(doc, session_num, session_data, head, secretary):
    """Добавляет таблицу с докладами для заседания"""
    # Заголовок заседания
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Заседание {session_num}")
    run.bold = True
    run.font.size = Pt(10)

    # Дата и место
    first_item = session_data[0]
    date_str = format_date_russian(first_item["date"])
    time = first_item["time"]
    aud = first_item["aud"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{date_str}, {time}, ул. Б. Морская, д. 67, ауд. {aud}")
    run.font.size = Pt(10)

    # Руководство
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Научный руководитель секции – {head}Секретарь – {secretary}")
    run.font.size = Pt(10)

    # Создаем таблицу
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.allow_autofit = False

    # Ширины столбцов
    column_widths = [Cm(1.06), Cm(8.82), Cm(2.47), Cm(3.53)]

    # Установка ширины таблицы и столбцов
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Удаляем существующие gridCol
    tblGrid = OxmlElement('w:tblGrid')
    tbl.insert(0, tblGrid)

    # Добавляем новые gridCol
    for width in column_widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(width.cm * 567)))
        tblGrid.append(gridCol)

    # Устанавливаем границы
    set_table_borders(table)

    # Заголовки таблицы
    headers = [
        "№ п/п",
        "Фамилия и инициалы докладчика, название доклада",
        "Статус(магистр/студент)",
        "Решение"
    ]

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_format(hdr_cells[i])
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Добавление докладов
    for i, item in enumerate(session_data, 1):
        if not item["full_name"] or not item["topic"]:
            continue

        row = table.add_row()
        row_cells = row.cells

        # Номер
        row_cells[0].text = str(i)
        set_cell_format(row_cells[0])

        # ФИО и тема
        p = row_cells[1].paragraphs[0]
        run = p.add_run(f"{item['full_name']}. {item['topic']}")
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Статус
        status_text = ""
        if item["status"] and item["group"]:
            status_text = f"{item['status']}\nгр. {item['group']}"
        elif item["group"]:
            status_text = f"гр. {item['group']}"
        elif item["status"]:
            status_text = item["status"]

        row_cells[2].text = status_text
        set_cell_format(row_cells[2])

        # Решение
        row_cells[3].text = ""
        set_cell_format(row_cells[3])


def generate_conference_report(conference_data, contributions, output_path):
    """Генерирует файл отчета конференции"""
    doc = Document()

    # Установка полей документа
    section = doc.sections[0]
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    set_document_styles(doc)

    # Заголовок отчета
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Отчет о проведении {conference_data['number']} МСНК ГУАП")
    run.bold = True
    run.font.size = Pt(10)

    # Секция кафедры
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Секция 43. Кафедра компьютерных технологий и программной инженерии")
    run.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

    # Обработка данных
    processed_entries = process_contributions(contributions)

    # Группировка по заседаниям
    sessions = {}
    for entry in processed_entries:
        session = entry["session"]
        if session not in sessions:
            sessions[session] = []
        sessions[session].append(entry)

    # Сортировка заседаний по дате
    def get_session_date(session_data):
        return session_data[0]["date"] if session_data else ""

    sorted_sessions = sorted(
        sessions.items(),
        key=lambda x: get_session_date(x[1])
    )

    # Добавление заседаний
    for session_num, session_data in sorted_sessions:
        add_session_table(
            doc,
            session_num,
            session_data,
            conference_data["head"],
            conference_data["secretary"]
        )
        doc.add_paragraph()  # Пустая строка между заседаниями

    # Подпись руководителя
    doc.add_paragraph("Научный руководитель секции                                    ___________________ / " +
                      conference_data["head"].split(",")[0])

    doc.save(output_path)