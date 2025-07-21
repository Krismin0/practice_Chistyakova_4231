import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter import messagebox
from docx.shared import Cm


def create_accepted_papers_list(input_json_path, output_docx_path, leader_name, leader_email, leader_phone, conf_number):
    try:
        # Загрузка данных из JSON
        with open(input_json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Фильтрация принятых работ
        accepted_papers = [paper for paper in data['papers'] if paper['state']['name'] == 'accepted']

        if not accepted_papers:
            messagebox.showwarning("Предупреждение", "Нет принятых работ для формирования списка.")
            return

        # Создание документа Word
        doc = Document()
        # Установка полей документа
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        style.paragraph_format.line_spacing = 1.15

        # Заголовок
        title = doc.add_paragraph('Список представляемых к публикации докладов')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title.runs[0]
        run.font.bold = True
        run.font.italic = True
        title.paragraph_format.space_after = Pt(0)

        p = doc.add_paragraph("")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15


        # Информация о кафедре и руководителе
        def add_info_paragraph(text):
            p = doc.add_paragraph(text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.left_indent = Cm(2)
            run = p.runs[0]
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(0)

        add_info_paragraph("Кафедра № 43 компьютерных технологий и программной инженерии")
        add_info_paragraph(leader_name)
        add_info_paragraph(f"e-mail: {leader_email}")
        add_info_paragraph(f"тел.: {leader_phone}")

        p = doc.add_paragraph("")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15


        for idx, paper in enumerate(accepted_papers, start=1):
            submitter = paper['revisions'][0]['submitter']
            full_name = submitter['full_name'].strip()
            parts = full_name.split()

            if len(parts) >= 2:
                last_name = parts[0]
                initials = f"{parts[1][0]}."
                if len(parts) >= 3:
                    initials += f"{parts[2][0]}."
                author_name = f"{last_name} {initials}"
            else:
                author_name = full_name

            paper_title = paper['contribution']['title']

            # Добавляем строку с табуляцией
            p = doc.add_paragraph(f"\t{idx}.\t{author_name}, {paper_title}")
            p.paragraph_format.left_indent = Cm(0)  # весь абзац без отступа
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            run = p.runs[0]

        # Добавляем подпись
        doc.add_paragraph("\n")
        signature = doc.add_paragraph(f"Руководитель УНИДС                                               {leader_name} ")
        signature.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        signature.paragraph_format.left_indent = Cm(2)
        run = signature.runs[0]
        run.font.size = Pt(12)

        # Сохранение документа
        doc.save(output_docx_path)

    except Exception as e:
        messagebox.showerror("Ошибка", f"Произошла ошибка при создании документа:\n{e}")